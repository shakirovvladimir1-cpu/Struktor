import json
import logging
import re

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from google import genai
from google.genai import types

from prompts import SYSTEM_PROMPT

logger = logging.getLogger(__name__)

GEMINI_MODEL = "gemini-2.5-flash-lite"
BATCH_SIZE = 20   # paragraphs per Gemini call
MAX_CHARS_PER_BATCH = 20000  # max total text chars per Gemini call


def iter_paragraphs(doc: Document):
    """Yield all paragraphs from document body and table cells, in order."""
    for block in doc.element.body:
        if block.tag == qn("w:p"):
            yield Paragraph(block, doc)
        elif block.tag == qn("w:tbl"):
            for p_elem in block.iter(qn("w:p")):
                yield Paragraph(p_elem, doc)


def split_para_by_linebreaks(para: Paragraph) -> list[str]:
    """Split paragraph text by <w:br/> line breaks, return non-empty chunks."""
    chunks = []
    current = []
    for elem in para._element.iter():
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag == "t" and elem.text:
            current.append(elem.text)
        elif tag == "br":
            text = "".join(current).strip()
            if text:
                chunks.append(text)
            current = []
    text = "".join(current).strip()
    if text:
        chunks.append(text)
    return chunks if chunks else ([para.text.strip()] if para.text.strip() else [])


CHUNK_SIZE = 5000  # max chars per virtual chunk when splitting giant paragraphs


def split_text_into_chunks(text: str, max_chars: int = CHUNK_SIZE) -> list[str]:
    """Split a large text into chunks of up to max_chars, breaking at sentence/comma boundaries."""
    if len(text) <= max_chars:
        return [text]
    chunks = []
    while text:
        if len(text) <= max_chars:
            chunks.append(text)
            break
        # Find best break point (comma, period, semicolon) near max_chars
        candidate = text[:max_chars]
        for sep in (",", ";", ".", " "):
            idx = candidate.rfind(sep)
            if idx > max_chars // 2:
                candidate = text[:idx + 1]
                break
        chunks.append(candidate.strip())
        text = text[len(candidate):].strip()
    return [c for c in chunks if c]


def collect_paragraphs(doc: Document) -> list[dict]:
    """Return list of {id, para, text} for all non-empty paragraphs.
    Giant single-paragraph documents are split into logical chunks."""
    raw = []
    for para in iter_paragraphs(doc):
        text = para.text.strip()
        if not text:
            continue
        # Pre-process before chunking
        text = fix_latin_chars(text)
        text = preprocess_ranges(text)
        if len(text) > CHUNK_SIZE:
            # Try line-break split first
            lines = split_para_by_linebreaks(para)
            if len(lines) > 1:
                for line in lines:
                    raw.append({"para": para, "text": line, "virtual": True})
                continue
            # Fall back to character-based chunking
            chunks = split_text_into_chunks(text)
            if len(chunks) > 1:
                for chunk in chunks:
                    raw.append({"para": para, "text": chunk, "virtual": True})
                continue
        raw.append({"para": para, "text": text, "virtual": False})

    result = [{"id": i, **r} for i, r in enumerate(raw)]
    red_count = sum(1 for r in result if '[[RED:' in r['text'])
    logger.info(f"Ranges found (red markers): {red_count}")
    return result


RED_MARKER = re.compile(r'\[\[RED:(.*?)\]\]')
RED_COLOR = RGBColor(0xFF, 0x00, 0x00)


def apply_text_to_para(para: Paragraph, new_text: str):
    """Replace paragraph text. Handles [[RED:...]] markers for red coloring."""
    # Clear all existing runs
    p_elem = para._element
    for r in p_elem.findall(qn("w:r")):
        p_elem.remove(r)

    parts = RED_MARKER.split(new_text)
    # split result: [normal, red, normal, red, ...] — odd indices are red
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        if i % 2 == 1:
            run.font.color.rgb = RED_COLOR


def make_paragraph_element(text: str):
    """Create a minimal w:p XML element with given text (no RED marker support)."""
    new_p = OxmlElement("w:p")
    new_r = OxmlElement("w:r")
    new_t = OxmlElement("w:t")
    new_t.text = text
    new_r.append(new_t)
    new_p.append(new_r)
    return new_p


def make_paragraph_with_markers(doc: Document, text: str) -> Paragraph:
    """Create a new Paragraph object supporting [[RED:...]] markers."""
    para = doc.add_paragraph()
    # Remove the paragraph from doc body — caller will insert it elsewhere
    doc.element.body.remove(para._element)
    apply_text_to_para(para, text)
    return para


def insert_paragraphs_before(para: Paragraph, texts: list[str], doc: Document):
    """Insert multiple paragraphs immediately before `para`, in order."""
    ref_elem = para._element
    for text in texts:
        if RED_MARKER.search(text):
            new_para = make_paragraph_with_markers(doc, text)
            ref_elem.addprevious(new_para._element)
        else:
            new_p = make_paragraph_element(text)
            ref_elem.addprevious(new_p)


_LATIN_TO_CYR = {
    'a': 'а', 'A': 'А', 'e': 'е', 'E': 'Е', 'o': 'о', 'O': 'О',
    'p': 'р', 'P': 'Р', 'c': 'с', 'C': 'С', 'x': 'х', 'X': 'Х',
    'y': 'у', 'Y': 'У', 'B': 'В', 'H': 'Н', 'M': 'М', 'T': 'Т',
    'K': 'К', 'u': 'и', 'l': 'л',
}


def fix_latin_chars(text: str) -> str:
    """Replace Latin lookalikes with Cyrillic only in mixed-script words."""
    def fix_word(m):
        word = m.group()
        if re.search(r'[а-яёА-ЯЁ]', word) and re.search(r'[a-zA-Z]', word):
            return ''.join(_LATIN_TO_CYR.get(ch, ch) for ch in word)
        return word
    return re.sub(r'\S+', fix_word, text)


_STOP_PATTERNS = [
    r'\bдолжн[аоы]?\s+быть\b',
    r'\bдолжен\s+быть\b',
    r'\bдолжн[аоы]?\b',
    r'\bдолжен\b',
    r'\bне\s+менее\s+чем\b',
    r'\bне\s+более\s+чем\b',
    r'\bне\s+менее\b',
    r'\bне\s+более\b',
    r'\bне\s+(?:хуже|ниже|выше|меньше|больше)(?:\s+чем)?\b',
    r'\bнеменее\b',
    r'\bминимум\b',
    r'\bмаксимум\b',
    r'\bпримерно\b',
    r'\bориентировочно\b',
    r'\bприблизительно\b',
    r'\bоколо\b',
    r'\bв\s+среднем\b',
    r'\bгде-то\b',
    r'\bна\s+примере\b',
    r'\bнапример\b',
    r'\bможет\s+быть\b',
    r'\bлибо\b',
    r'\bне\s+превышать\b',
    r'\bне\s+превышает\b',
    r'\bв\s+районе\b',
]


def postprocess_text(text: str) -> str:
    """Post-process after Gemini: более/менее → ±1, remove leftover stop words."""
    # Step 0: "до X[unit]" (standalone, = не более X) → keep X
    text = re.sub(r'\bдо\s+([\d,\.]+)', r'\1', text, flags=re.IGNORECASE)

    # Step 1: remove "не + qualifier" → keep the number
    text = re.sub(
        r'\bне\s+(?:более|менее|больше|меньше|выше|ниже|хуже|лучше)(?:\s+чем)?\s+',
        '', text, flags=re.IGNORECASE
    )

    # Step 2: "более/выше/больше/старше/свыше X" → X + step
    def bolee(m):
        try:
            val = float(m.group(1).replace(',', '.'))
            step = 0.1 if val <= 1 else 1
            new_val = round(val + step, 10)
            result = str(int(new_val)) if new_val == int(new_val) else str(round(new_val, 4)).rstrip('0').rstrip('.')
            unit = m.group(2) or ''
            return result + (' ' if unit else '') + unit
        except ValueError:
            return m.group(0)
    text = re.sub(r'\b(?:более|выше|больше|старше|свыше)\s+([\d,\.]+)\s*([^\s,;\.\[]{0,8})',
                  bolee, text, flags=re.IGNORECASE)

    # Step 3: "менее/ниже/меньше/младше/хуже/тоньше X" → X - step
    def menee(m):
        try:
            val = float(m.group(1).replace(',', '.'))
            step = 0.1 if val <= 1 else 1
            new_val = round(val - step, 10)
            result = str(int(new_val)) if new_val == int(new_val) else str(round(new_val, 4)).rstrip('0').rstrip('.')
            unit = m.group(2) or ''
            return result + (' ' if unit else '') + unit
        except ValueError:
            return m.group(0)
    text = re.sub(r'\b(?:менее|ниже|меньше|младше|хуже|тоньше)\s+([\d,\.]+)\s*([^\s,;\.\[]{0,8})',
                  menee, text, flags=re.IGNORECASE)

    # Step 4: символы сравнения → числа ±1
    # "> X" или "> X" → X+1
    def gt(m):
        try:
            val = float(m.group(1).replace(',', '.'))
            unit = m.group(2) or ''
            return str(int(val) + 1) + (' ' if unit else '') + unit
        except ValueError:
            return m.group(0)
    text = re.sub(r'[>]\s*([\d,\.]+)\s*([^\s,;\.\[]{0,8})', gt, text)

    # "< X" → X-1
    def lt(m):
        try:
            val = float(m.group(1).replace(',', '.'))
            unit = m.group(2) or ''
            return str(int(val) - 1) + (' ' if unit else '') + unit
        except ValueError:
            return m.group(0)
    text = re.sub(r'[<]\s*([\d,\.]+)\s*([^\s,;\.\[]{0,8})', lt, text)

    # ">= X" или "≥ X" → оставить X (убрать символ)
    text = re.sub(r'[≥]\s*', '', text)
    text = re.sub(r'>=\s*', '', text)
    # "<= X" или "≤ X" → оставить X
    text = re.sub(r'[≤]\s*', '', text)
    text = re.sub(r'<=\s*', '', text)

    # Step 5: "или" → "и"
    text = re.sub(r'\bили\b', 'и', text, flags=re.IGNORECASE)

    # Step 6: remove ±
    text = text.replace('±', '')

    # Step 7: remove remaining stop words
    for pattern in _STOP_PATTERNS:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)

    # Cleanup artifacts
    text = re.sub(r'\s{2,}', ' ', text)
    text = re.sub(r'\s+([,;:\.])', r'\1', text)
    return text.strip()


def preprocess_ranges(text: str) -> str:
    """Replace numeric ranges with [[RED:median]] before sending to Gemini."""
    # Pattern: "не менее X[unit], но не более Y[unit]" or "не менее X[unit] и не более Y[unit]"
    def replace_range(m):
        x_str = m.group(1).replace(',', '.')
        y_str = m.group(3).replace(',', '.')
        unit = m.group(4) or ''
        try:
            x, y = float(x_str), float(y_str)
            median = round((x + y) / 2)
            return f'[[RED:{median}]]{unit}'
        except ValueError:
            return m.group(0)

    # "не менее X[unit], но не более Y[unit]"
    text = re.sub(
        r'не\s+менее\s+([\d,\.]+)\s*([^,]{0,10}?),?\s+(?:но\s+)?не\s+более\s+([\d,\.]+)\s*([^\s,;\.]{0,10})',
        replace_range, text, flags=re.IGNORECASE
    )
    # "от X до Y[unit]"
    def replace_from_to(m):
        x_str = m.group(1).replace(',', '.')
        y_str = m.group(2).replace(',', '.')
        unit = m.group(3) or ''
        try:
            x, y = float(x_str), float(y_str)
            median = round((x + y) / 2)
            return f'[[RED:{median}]]{unit}'
        except ValueError:
            return m.group(0)

    text = re.sub(
        r'от\s+([\d,\.]+)\s*[^\d\s,;]{0,10}\s+до\s+([\d,\.]+)\s*([^\s,;\.]{0,10})',
        replace_from_to, text, flags=re.IGNORECASE
    )
    return text


def fix_json_escapes(text: str) -> str:
    """Fix invalid backslash escapes inside JSON string values."""
    # Valid JSON escape chars: " \ / b f n r t u
    valid = set('"\\\/bfnrtu')
    result = []
    i = 0
    while i < len(text):
        ch = text[i]
        if ch == '\\' and i + 1 < len(text):
            next_ch = text[i + 1]
            if next_ch in valid:
                result.append(ch)
                result.append(next_ch)
                i += 2
            else:
                # Invalid escape — double the backslash
                result.append('\\\\')
                i += 1
        else:
            result.append(ch)
            i += 1
    return "".join(result)


def parse_gemini_response(raw: str) -> list[dict]:
    """Extract JSON array from Gemini response, stripping markdown fences."""
    text = raw.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    text = text.strip()
    # Fix trailing commas before ] or } (not valid JSON but Gemini sometimes adds them)
    text = re.sub(r",\s*([\]\}])", r"\1", text)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        text = fix_json_escapes(text)
        return json.loads(text)


def call_gemini(client: genai.Client, batch: list[dict]) -> dict[int, list[str]]:
    """Send a batch to Gemini, return {id: [list_of_paragraphs]} mapping."""
    input_data = [{"id": p["id"], "text": p["text"]} for p in batch]
    prompt = (
        SYSTEM_PROMPT
        + "\n\nВходные данные:\n"
        + json.dumps(input_data, ensure_ascii=False, indent=2)
    )
    response = client.models.generate_content(
        model=GEMINI_MODEL,
        contents=prompt,
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
        ),
    )
    cleaned = parse_gemini_response(response.text)

    result = {}
    for item in cleaned:
        paragraphs = item.get("paragraphs")
        if paragraphs is None:
            paragraphs = [item.get("text", "")]
        result[item["id"]] = [postprocess_text(p) for p in paragraphs]
    return result


def process_docx(input_path: str, output_path: str, gemini_api_key: str):
    """
    Read DOCX, clean and restructure all paragraphs with Gemini 2.5 Flash,
    save result to output_path.

    For item paragraphs (numbered goods), inserts 5 structured header lines
    (Наименование, Страна, Завод, Год, Гарантийный срок) before the spec text.
    """
    client = genai.Client(api_key=gemini_api_key)

    doc = Document(input_path)
    paragraphs = collect_paragraphs(doc)

    if not paragraphs:
        raise ValueError("Документ пустой или не содержит текстовых абзацев.")

    logger.info(f"Total paragraphs to process: {len(paragraphs)}")

    # Process in batches, respecting both paragraph count and char limits
    cleaned_map: dict[int, list[str]] = {}
    batch: list[dict] = []
    batch_chars = 0
    batch_num = 0

    def flush_batch(b):
        nonlocal batch_num
        if not b:
            return
        batch_num += 1
        result = call_gemini(client, b)
        cleaned_map.update(result)
        logger.info(f"Batch {batch_num} done ({len(b)} paragraphs)")

    for p in paragraphs:
        p_len = len(p["text"])
        # If single paragraph exceeds limit — send it alone
        if p_len > MAX_CHARS_PER_BATCH:
            flush_batch(batch)
            batch, batch_chars = [], 0
            flush_batch([p])
            continue
        # If adding this paragraph would exceed limits — flush first
        if batch and (len(batch) >= BATCH_SIZE or batch_chars + p_len > MAX_CHARS_PER_BATCH):
            flush_batch(batch)
            batch, batch_chars = [], 0
        batch.append(p)
        batch_chars += p_len

    flush_batch(batch)

    has_virtual = any(p.get("virtual") for p in paragraphs)

    if has_virtual:
        # Document uses line-breaks instead of paragraphs — build new doc from scratch
        out_doc = Document()
        for p in paragraphs:
            lines = cleaned_map.get(p["id"], [p["text"]])
            for line in lines:
                new_para = out_doc.add_paragraph()
                apply_text_to_para(new_para, line)
        out_doc.save(output_path)
        logger.info(f"Virtual-paragraph doc: wrote {len(paragraphs)} lines as real paragraphs")
        return

    # Normal doc: Apply changes in-place
    changed = 0
    inserted = 0
    for p in paragraphs:
        new_paragraphs = cleaned_map.get(p["id"])
        if not new_paragraphs:
            continue

        if len(new_paragraphs) == 1:
            new_text = new_paragraphs[0]
            if new_text != p["text"]:
                apply_text_to_para(p["para"], new_text)
                changed += 1
        else:
            headers = new_paragraphs[:-1]
            spec_text = new_paragraphs[-1]
            insert_paragraphs_before(p["para"], headers, doc)
            apply_text_to_para(p["para"], spec_text)
            inserted += len(headers)
            changed += 1

    logger.info(
        f"Changed {changed} paragraphs, inserted {inserted} new header paragraphs"
    )
    doc.save(output_path)
